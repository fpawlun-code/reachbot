"""
Moduł eksportu danych do różnych formatów.

Obsługiwane formaty:
- CSV
- Excel (XLSX)
- Word (DOCX)
- JSON
"""
import csv
import json
import logging
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

import sys
sys.path.insert(0, str(__file__).rsplit("/", 2)[0])
from config import OUTPUT_DIR

logger = logging.getLogger(__name__)


class DataExporter:
    """
    Eksportuje dane o firmach do różnych formatów.
    """

    # Kolumny do eksportu
    COLUMNS = [
        ("name", "Nazwa firmy"),
        ("industry", "Branża"),
        ("address", "Adres"),
        ("phone", "Telefon"),
        ("email", "Email"),
        ("facebook", "Facebook"),
        ("instagram", "Instagram"),
        ("linkedin", "LinkedIn"),
        ("website", "Strona WWW"),
        ("has_website", "Ma stronę?"),
        ("source", "Źródło"),
    ]

    def __init__(self, output_dir: Path = None):
        """
        Inicjalizuje eksporter.

        Args:
            output_dir: Katalog wyjściowy (domyślnie: output/)
        """
        self.output_dir = output_dir or OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        businesses: List[Dict],
        filename: str = None,
        format: str = "xlsx"
    ) -> Path:
        """
        Eksportuje dane do wybranego formatu.

        Args:
            businesses: Lista firm do eksportu
            filename: Nazwa pliku (bez rozszerzenia)
            format: Format wyjściowy (csv, xlsx, docx, json)

        Returns:
            Ścieżka do utworzonego pliku
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"firmy_szczecin_{timestamp}"

        format = format.lower()
        exporters = {
            "csv": self._export_csv,
            "xlsx": self._export_xlsx,
            "excel": self._export_xlsx,
            "docx": self._export_docx,
            "word": self._export_docx,
            "json": self._export_json,
        }

        if format not in exporters:
            raise ValueError(f"Nieobsługiwany format: {format}. Dostępne: {list(exporters.keys())}")

        return exporters[format](businesses, filename)

    def _prepare_data(self, businesses: List[Dict]) -> List[Dict]:
        """Przygotowuje dane do eksportu."""
        prepared = []
        for biz in businesses:
            row = {}
            for key, label in self.COLUMNS:
                value = biz.get(key, "")
                # Konwertuj boolean na tekst
                if isinstance(value, bool):
                    value = "Tak" if value else "Nie"
                row[label] = value or ""
            prepared.append(row)
        return prepared

    def _export_csv(self, businesses: List[Dict], filename: str) -> Path:
        """Eksportuje do CSV."""
        filepath = self.output_dir / f"{filename}.csv"
        data = self._prepare_data(businesses)

        if not data:
            logger.warning("No data to export")
            return filepath

        headers = [label for _, label in self.COLUMNS]

        with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=headers, delimiter=";")
            writer.writeheader()
            writer.writerows(data)

        logger.info(f"Exported {len(businesses)} businesses to {filepath}")
        return filepath

    def _export_xlsx(self, businesses: List[Dict], filename: str) -> Path:
        """Eksportuje do Excel."""
        filepath = self.output_dir / f"{filename}.xlsx"
        data = self._prepare_data(businesses)

        if not data:
            logger.warning("No data to export")
            # Utwórz pusty plik
            pd.DataFrame().to_excel(filepath, index=False)
            return filepath

        df = pd.DataFrame(data)

        # Zapisz z formatowaniem
        with pd.ExcelWriter(filepath, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Firmy", index=False)

            # Formatowanie
            worksheet = writer.sheets["Firmy"]

            # Szerokość kolumn
            column_widths = {
                "Nazwa firmy": 30,
                "Branża": 20,
                "Adres": 40,
                "Telefon": 15,
                "Email": 30,
                "Facebook": 35,
                "Instagram": 30,
                "LinkedIn": 35,
                "Strona WWW": 35,
                "Ma stronę?": 12,
                "Źródło": 15,
            }

            for idx, col in enumerate(df.columns):
                width = column_widths.get(col, 20)
                worksheet.column_dimensions[chr(65 + idx)].width = width

        logger.info(f"Exported {len(businesses)} businesses to {filepath}")
        return filepath

    def _export_docx(self, businesses: List[Dict], filename: str) -> Path:
        """Eksportuje do dokumentu Word."""
        filepath = self.output_dir / f"{filename}.docx"
        data = self._prepare_data(businesses)

        doc = Document()

        # Tytuł
        title = doc.add_heading("Raport - Firmy bez stron internetowych", 0)

        # Informacje o raporcie
        doc.add_paragraph(f"Data wygenerowania: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Liczba firm: {len(businesses)}")
        doc.add_paragraph(f"Lokalizacja: Szczecin")
        doc.add_paragraph()

        if not data:
            doc.add_paragraph("Brak danych do wyświetlenia.")
            doc.save(filepath)
            return filepath

        # Grupuj po branżach
        by_industry = {}
        for biz in data:
            industry = biz.get("Branża", "Inne")
            if industry not in by_industry:
                by_industry[industry] = []
            by_industry[industry].append(biz)

        # Dla każdej branży stwórz sekcję
        for industry, biz_list in sorted(by_industry.items()):
            doc.add_heading(f"{industry} ({len(biz_list)})", level=1)

            # Tabela dla branży
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Nagłówki
            headers = ["Nazwa", "Adres", "Telefon", "Email", "Social Media"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Pogrubienie nagłówków
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Dane
            for biz in biz_list:
                row = table.add_row().cells

                row[0].text = biz.get("Nazwa firmy", "")
                row[1].text = biz.get("Adres", "")
                row[2].text = biz.get("Telefon", "")
                row[3].text = biz.get("Email", "")

                # Social media w jednej kolumnie
                social = []
                if biz.get("Facebook"):
                    social.append(f"FB: {biz['Facebook']}")
                if biz.get("Instagram"):
                    social.append(f"IG: {biz['Instagram']}")
                if biz.get("LinkedIn"):
                    social.append(f"LI: {biz['LinkedIn']}")
                row[4].text = "\n".join(social)

            doc.add_paragraph()  # Odstęp między sekcjami

        doc.save(filepath)
        logger.info(f"Exported {len(businesses)} businesses to {filepath}")
        return filepath

    def _export_json(self, businesses: List[Dict], filename: str) -> Path:
        """Eksportuje do JSON."""
        filepath = self.output_dir / f"{filename}.json"

        output = {
            "generated_at": datetime.now().isoformat(),
            "total_count": len(businesses),
            "location": "Szczecin",
            "businesses": businesses,
        }

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(output, f, ensure_ascii=False, indent=2)

        logger.info(f"Exported {len(businesses)} businesses to {filepath}")
        return filepath

    def export_summary(self, businesses: List[Dict], filename: str = None) -> Path:
        """
        Eksportuje podsumowanie statystyczne.

        Args:
            businesses: Lista firm
            filename: Nazwa pliku

        Returns:
            Ścieżka do pliku
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"podsumowanie_{timestamp}"

        filepath = self.output_dir / f"{filename}.txt"

        # Statystyki
        total = len(businesses)
        with_email = sum(1 for b in businesses if b.get("email"))
        with_phone = sum(1 for b in businesses if b.get("phone"))
        with_facebook = sum(1 for b in businesses if b.get("facebook"))
        with_instagram = sum(1 for b in businesses if b.get("instagram"))
        without_website = sum(1 for b in businesses if not b.get("has_website"))

        # Grupuj po branżach
        by_industry = {}
        for b in businesses:
            ind = b.get("industry", "Inne")
            by_industry[ind] = by_industry.get(ind, 0) + 1

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("=" * 50 + "\n")
            f.write("PODSUMOWANIE SKANOWANIA FIRM - SZCZECIN\n")
            f.write("=" * 50 + "\n\n")

            f.write(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

            f.write("STATYSTYKI OGÓLNE:\n")
            f.write("-" * 30 + "\n")
            f.write(f"Łączna liczba firm: {total}\n")
            f.write(f"Firmy BEZ strony www: {without_website} ({without_website/total*100:.1f}%)\n")
            f.write(f"Firmy z emailem: {with_email} ({with_email/total*100:.1f}%)\n")
            f.write(f"Firmy z telefonem: {with_phone} ({with_phone/total*100:.1f}%)\n")
            f.write(f"Firmy z Facebook: {with_facebook} ({with_facebook/total*100:.1f}%)\n")
            f.write(f"Firmy z Instagram: {with_instagram} ({with_instagram/total*100:.1f}%)\n\n")

            f.write("PODZIAŁ NA BRANŻE:\n")
            f.write("-" * 30 + "\n")
            for industry, count in sorted(by_industry.items(), key=lambda x: -x[1]):
                f.write(f"  {industry}: {count}\n")

        logger.info(f"Summary exported to {filepath}")
        return filepath
